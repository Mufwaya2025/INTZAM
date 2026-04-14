#!/usr/bin/env python3
"""
generate_handover_doc.py
IntZam Microfinance Ltd — LMS–Odoo IT Handover Document Generator
Run: python generate_handover_doc.py
Output: /mnt/f/LMS/IntZam_LMS_Odoo_IT_Handover.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUTPUT_PATH = "/mnt/f/LMS/IntZam_LMS_Odoo_IT_Handover.docx"
PREPARED_DATE = "17 March 2026"
COMPANY = "IntZam Microfinance Ltd"
RECIPIENT = "IT Manager"
AUTHOR = "[PENDING — Author name]"

# ── Colour palette ────────────────────────────────────────────────────────────
BRAND_BLUE   = RGBColor(0x1A, 0x3C, 0x6E)   # dark navy
ACCENT_BLUE  = RGBColor(0x1E, 0x6F, 0xBB)   # mid blue
LIGHT_BLUE   = RGBColor(0xD6, 0xE4, 0xF5)   # table header fill
LIGHT_GREY   = RGBColor(0xF4, 0xF6, 0xF8)   # alt row fill
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
BLACK        = RGBColor(0x1A, 0x1A, 0x1A)
RED_PENDING  = RGBColor(0xC0, 0x39, 0x2B)
GREEN_DONE   = RGBColor(0x1E, 0x8B, 0x4C)
AMBER        = RGBColor(0xE6, 0x7E, 0x22)


# ═══════════════════════════════════════════════════════════════════════════════
# Helper utilities
# ═══════════════════════════════════════════════════════════════════════════════

def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_col_widths(table, widths_cm):
    """Set explicit column widths (list of cm values)."""
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


def cell_para(cell, text, bold=False, italic=False, font_size=9,
              color=None, align=WD_ALIGN_PARAGRAPH.LEFT, wrap=True):
    """Write text into a table cell paragraph, return the run."""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    run.font.color.rgb = color or BLACK
    return run


def add_header_row(table, headers, col_widths=None, font_size=9):
    """Style the first row of a table as a dark-blue header."""
    hdr_row = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, '1A3C6E')
        cell_para(cell, hdr, bold=True, font_size=font_size,
                  color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    if col_widths:
        set_col_widths(table, col_widths)


def zebra_row(table, row_idx, data, alt=False, font_size=9, col_colors=None):
    """Fill a data row; alternate grey rows."""
    row = table.rows[row_idx]
    bg  = 'F4F6F8' if alt else 'FFFFFF'
    for i, val in enumerate(data):
        cell  = row.cells[i]
        color = col_colors[i] if col_colors and i < len(col_colors) else BLACK
        set_cell_bg(cell, bg)
        cell_para(cell, str(val), font_size=font_size, color=color)


def add_table(doc, headers, rows, col_widths=None, font_size=9, col_colors=None):
    """Create a complete styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_header_row(table, headers, col_widths, font_size)
    for i, row_data in enumerate(rows):
        zebra_row(table, i + 1, row_data, alt=(i % 2 == 1),
                  font_size=font_size, col_colors=col_colors)
    return table


def heading(doc, text, level, color=None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color or BRAND_BLUE
    return p


def body(doc, text, bold=False, italic=False, color=None, size=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color or BLACK
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def pending(doc, note):
    p = doc.add_paragraph()
    run = p.add_run(f"⚠  PENDING: {note}")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RED_PENDING
    return p


def section_rule(doc):
    """Thin horizontal rule between sections."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A3C6E')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def page_break(doc):
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# Cover page
# ═══════════════════════════════════════════════════════════════════════════════

def build_cover(doc):
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(COMPANY)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = BRAND_BLUE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("LMS — Odoo 17 Integration")
    r2.bold = True
    r2.font.size = Pt(16)
    r2.font.color.rgb = ACCENT_BLUE

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("IT HANDOVER DOCUMENT")
    r3.bold = True
    r3.font.size = Pt(14)
    r3.font.color.rgb = BLACK

    doc.add_paragraph()
    doc.add_paragraph()

    meta = [
        ("Document version", "1.0"),
        ("Prepared date",    PREPARED_DATE),
        ("Prepared by",      AUTHOR),
        ("Handover recipient", RECIPIENT),
        ("Classification",   "CONFIDENTIAL — Internal Use Only"),
    ]
    tbl = doc.add_table(rows=len(meta), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(meta):
        row = tbl.rows[i]
        set_cell_bg(row.cells[0], 'D6E4F5')
        cell_para(row.cells[0], k, bold=True, font_size=10, color=BRAND_BLUE)
        set_cell_bg(row.cells[1], 'FFFFFF')
        color = RED_PENDING if v.startswith('[PENDING') else BLACK
        cell_para(row.cells[1], v, font_size=10, color=color)
    set_col_widths(tbl, [5.5, 9.5])

    page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# Section 1 — System Overview
# ═══════════════════════════════════════════════════════════════════════════════

def build_section1(doc):
    heading(doc, "SECTION 1 — System Overview", 1)
    section_rule(doc)

    body(doc, (
        "IntZam Microfinance Ltd operates a bespoke Loan Management System (LMS) that manages "
        "the full loan lifecycle — from application through disbursement, repayment collections, "
        "IFRS 9 provisioning, and write-off / recovery. This document covers the integration "
        "between the LMS and Odoo 17 Community, which serves as the organisation's accounting "
        "general ledger and financial reporting platform."
    ))

    heading(doc, "1.1  Technology Stack", 2)
    rows = [
        ("LMS Backend",         "Django 5.0 + Django REST Framework, Python 3.12"),
        ("LMS Admin Frontend",  "React 19 + TypeScript + Vite  (port 5173)"),
        ("LMS Client PWA",      "React 19 + TypeScript + Vite  (port 5174)"),
        ("Accounting GL",       "Odoo 17 Community  (port 8069)"),
        ("Integration layer",   "Python XML-RPC via apps/accounting/odoo_client.py"),
        ("Database (LMS)",      "SQLite (development) / PostgreSQL 16 (production)"),
        ("Database (Odoo)",     "PostgreSQL 14 — database: odoo_lms_test (dev) / odoo_lms_production (prod)"),
        ("AI Risk Scoring",     "Google Gemini API  (apps/ai)"),
        ("Currency",            "ZMW (Zambian Kwacha) — regulatory: BoZ, ZRA, NAPSA"),
        ("Payment Gateway",     "Konse Konse (*543#) — sandbox integration built, live credentials pending"),
    ]
    add_table(doc, ["Component", "Detail"], rows, col_widths=[5, 10])

    doc.add_paragraph()
    heading(doc, "1.2  Integration Architecture Summary", 2)
    body(doc, (
        "The LMS communicates with Odoo exclusively via the Odoo XML-RPC API (endpoints "
        "/xmlrpc/2/common and /xmlrpc/2/object). Every financial event in the LMS — "
        "loan disbursement, repayment receipt, fee posting, ECL provisioning, write-off, "
        "and recovery — triggers a corresponding double-entry journal posting in Odoo, "
        "ensuring the accounting GL is always kept in sync with loan activity."
    ))
    body(doc, (
        "All Odoo calls are non-blocking: if Odoo is temporarily unavailable, the LMS "
        "transaction is committed to the LMS database first, and the error is logged. "
        "This prevents Odoo downtime from blocking loan operations. A dedicated service "
        "account (lms_bridge, UID 6) is used for all XML-RPC calls — the admin account "
        "is never used for automated integration."
    ))

    heading(doc, "1.3  Current Deployment Status", 2)
    rows2 = [
        ("Development / WSL",   "LIVE",    "Odoo running at localhost:8069 via ~/start_odoo.sh"),
        ("Production server",   "NOT YET PROVISIONED", "See Section 8 for deployment plan"),
        ("Konse Konse gateway", "SANDBOX CODE BUILT", "Awaiting live API credentials from Konse Konse"),
    ]
    colors = [None, None, None]
    tbl2 = doc.add_table(rows=1 + len(rows2), cols=3)
    tbl2.style = 'Table Grid'
    tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_header_row(tbl2, ["Environment", "Status", "Notes"], col_widths=[4, 4, 7])
    for i, (env, status, note) in enumerate(rows2):
        row = tbl2.rows[i + 1]
        bg = 'F4F6F8' if i % 2 == 1 else 'FFFFFF'
        set_cell_bg(row.cells[0], bg)
        set_cell_bg(row.cells[1], bg)
        set_cell_bg(row.cells[2], bg)
        cell_para(row.cells[0], env, font_size=9)
        s_color = GREEN_DONE if status == 'LIVE' else RED_PENDING if 'NOT' in status else AMBER
        cell_para(row.cells[1], status, bold=True, font_size=9, color=s_color)
        cell_para(row.cells[2], note, font_size=9)

    page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# Section 2 — Chart of Accounts
# ═══════════════════════════════════════════════════════════════════════════════

def build_section2(doc):
    heading(doc, "SECTION 2 — Chart of Accounts (LMS-Relevant Accounts)", 1)
    section_rule(doc)
    body(doc, (
        "The following accounts are those actively referenced by the LMS–Odoo integration. "
        "The full Odoo COA contains 146 accounts; this table shows the 35 accounts that are "
        "read or written by the LMS bridge. All accounts use ZMW as the operating currency "
        "and are IFRS 9 and BoZ-compliant. 'Custom' indicates the account was created "
        "specifically for the LMS; 'System' indicates it was present in the default Odoo COA."
    ))
    doc.add_paragraph()

    # Code | Name | Odoo Account Type | Internal Group | Journal(s) | Origin
    rows = [
        # Assets
        ("1101", "Bank — Main Account",                 "asset_cash",        "Asset",     "LRPY",            "System"),
        ("1102", "MTN MoMo Float",                      "asset_cash",        "Asset",     "—",               "Custom"),
        ("1103", "Airtel Money Float",                  "asset_cash",        "Asset",     "—",               "Custom"),
        ("1104", "Zamtel Kwacha Float",                 "asset_cash",        "Asset",     "—",               "Custom"),
        ("1105", "Bank — Main Clearing (ZMW)",          "asset_cash",        "Asset",     "LDIS / LRPY / LFEE", "Custom"),
        ("1107", "Konse Konse Gateway Float",           "asset_cash",        "Asset",     "LDIS",            "Custom"),
        ("1108", "Konse Konse Disbursement Clearing",   "asset_current",     "Asset",     "LDIS",            "Custom"),
        ("1111", "Loan Receivable — Stage 1",           "asset_receivable",  "Asset",     "LDIS / LRPY",     "Custom"),
        ("1112", "Loan Receivable — Stage 2",           "asset_receivable",  "Asset",     "LRPY",            "Custom"),
        ("1113", "Loan Receivable — Stage 3",           "asset_receivable",  "Asset",     "LRPY",            "Custom"),
        ("1120", "Interest Receivable",                 "asset_receivable",  "Asset",     "LINT",            "Custom"),
        ("1121", "Penalty Fees Receivable",             "asset_receivable",  "Asset",     "LFEE",            "Custom"),
        ("1130", "Loan Origination Fees Receivable",    "asset_current",     "Asset",     "LFEE",            "Custom"),
        ("1201", "Provision for Loan Losses — Stage 1", "asset_current",     "Asset",     "LPROV",           "Custom"),
        ("1202", "Provision for Loan Losses — Stage 2", "asset_current",     "Asset",     "LPROV",           "Custom"),
        ("1203", "Provision for Loan Losses — Stage 3", "asset_current",     "Asset",     "LPROV",           "Custom"),
        ("1204", "Write-off Reserve",                   "asset_current",     "Asset",     "LFEE",            "Custom"),
        ("1311", "Input VAT Recoverable",               "asset_current",     "Asset",     "—",               "Custom"),
        ("1321", "Deferred Origination Fees",           "asset_current",     "Asset",     "LFEE",            "Custom"),
        # Liabilities
        ("2105", "VAT Payable (ZRA)",                   "liability_current", "Liability", "LFEE",            "Custom"),
        ("2109", "Mobile Money Levy Payable",           "liability_current", "Liability", "—",               "Custom"),
        ("2111", "Konse Konse Fees Payable",            "liability_current", "Liability", "—",               "Custom"),
        # Income
        ("4101", "Interest Income — Stage 1 Loans",    "income",            "Income",    "LRPY / LINT",     "Custom"),
        ("4102", "Interest Income — Stage 2 Loans",    "income",            "Income",    "LRPY / LINT",     "Custom"),
        ("4103", "Interest Income — Stage 3 Loans",    "income",            "Income",    "LRPY / LINT",     "Custom"),
        ("4201", "Loan Origination Fees",               "income",            "Income",    "LFEE",            "Custom"),
        ("4202", "Penalty & Late Fees",                 "income",            "Income",    "LFEE",            "Custom"),
        ("4204", "Insurance Premium Income",            "income",            "Income",    "LFEE",            "Custom"),
        ("4302", "Recovery on Written-off Loans",       "income_other",      "Income",    "LRPY",            "Custom"),
        # Expenses
        ("5101", "Loan Loss Provision Expense — Stage 1", "expense",        "Expense",   "LPROV",           "Custom"),
        ("5102", "Loan Loss Provision Expense — Stage 2", "expense",        "Expense",   "LPROV",           "Custom"),
        ("5103", "Loan Loss Provision Expense — Stage 3", "expense",        "Expense",   "LPROV",           "Custom"),
        ("5104", "Loan Write-off Expense",               "expense",         "Expense",   "LFEE",            "Custom"),
        ("5223", "Konse Konse Gateway Fees",             "expense",         "Expense",   "LFEE",            "Custom"),
    ]
    add_table(doc,
              ["Code", "Account Name", "Odoo Type", "Group", "Journal(s)", "Origin"],
              rows,
              col_widths=[1.3, 5.5, 3.2, 1.8, 3.2, 1.5],
              font_size=8)

    doc.add_paragraph()
    body(doc, (
        "Note: Accounts 1102, 1103, 1104 (MoMo floats) are used for the Mobile Money Levy "
        "posting (post_momo_levy) which will be wired when the Konse Konse gateway is live. "
        "Account 1204 (Write-off Reserve) is debited at loan write-off in place of 5104 as "
        "the direct write-off expense — 5104 is retained on the COA for supplementary "
        "reporting purposes."
    ), italic=True, size=9)

    page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# Section 3 — Transaction Mapping
# ═══════════════════════════════════════════════════════════════════════════════

def build_section3(doc):
    heading(doc, "SECTION 3 — Transaction Mapping (LMS Event → Odoo Journal Entry)", 1)
    section_rule(doc)
    body(doc, (
        "Each table row describes one double-entry posting that the LMS bridge creates in Odoo. "
        "'Trigger' shows the Django view or management command that initiates the call, "
        "followed by the odoo_client.py method that executes the XML-RPC posting. "
        "Stage 1/2/3 refers to IFRS 9 loan classification (S1 = performing, S2 = "
        "significant credit risk, S3 = credit-impaired)."
    ))
    doc.add_paragraph()

    rows = [
        # Event | Dr Account | Cr Account | Amount Basis | Journal | Trigger | Python Method
        ("Loan Disbursement",
         "1111 Loan Receivable S1",
         "1105 Bank Clearing",
         "Full loan principal (loan.amount)",
         "LDIS",
         "DisburseLoanView → on_loan_disbursed()",
         "post_disbursement()"),

        ("Origination Fee",
         "1105 Bank Clearing",
         "4201 Origination Fees",
         "product.processing_fee% × loan.amount",
         "LFEE",
         "on_loan_disbursed()  (auto if processing_fee > 0)",
         "post_origination_fee()"),

        ("VAT on Origination Fee",
         "2105 VAT Payable (ZRA)",
         "4201 Origination Fees",
         "Origination fee × 16%",
         "LFEE",
         "on_loan_disbursed()  (alongside origination fee)",
         "post_vat_on_fees()"),

        ("Repayment — Principal",
         "1105 Bank Clearing",
         "1111/1112/1113 Loan Rec S1/S2/S3",
         "EIR principal split from amortisation schedule",
         "LRPY",
         "RepaymentView → on_payment_received()",
         "post_repayment()"),

        ("Repayment — Interest",
         "1105 Bank Clearing",
         "4101/4102/4103 Interest Income S1/S2/S3",
         "EIR interest split (IFRS 9 stage-aware)",
         "LRPY",
         "RepaymentView → on_payment_received()",
         "post_repayment()"),

        ("Penalty / Late Fee",
         "1105 Bank Clearing",
         "4202 Penalty & Late Fees",
         "penalty_amount parameter",
         "LFEE",
         "RepaymentView → on_payment_received()",
         "post_penalty()"),

        ("Interest Accrual (month-end)",
         "1120 Interest Receivable",
         "4101/4102/4103 Interest Income",
         "Accrued interest per loan (EIR, month-end)",
         "LINT",
         "sync_odoo_interest_accrual (cron, last day of month)",
         "post_interest_accrual()"),

        ("ECL Provision (monthly)",
         "5101/5102/5103 Provision Expense",
         "1201/1202/1203 Provision — S1/S2/S3",
         "ECL provision amount by IFRS 9 stage",
         "LPROV",
         "sync_odoo_ecl_provision (cron, 1st of month)",
         "post_ecl_provision()"),

        ("Stage 1 → Stage 2 Transfer",
         "5102 Provision Expense S2",
         "1202 Provision S2",
         "Provision amount for transferred loan",
         "LPROV",
         "update_loan_statuses → on_loan_stage_changed()",
         "post_stage_transfer()"),

        ("Stage 1/2 → Stage 3 Transfer",
         "5103 Provision Expense S3",
         "1203 Provision S3",
         "Provision amount for transferred loan",
         "LPROV",
         "update_loan_statuses → on_loan_stage_changed()",
         "post_stage_transfer()"),

        ("Loan Write-off",
         "1204 Write-off Reserve",
         "1113 Loan Receivable S3",
         "Outstanding loan balance at write-off",
         "LFEE",
         "WriteOffLoanView → on_loan_written_off()",
         "post_writeoff()"),

        ("Recovery Post Write-off",
         "1105 Bank Clearing",
         "4302 Recovery on Written-off Loans",
         "Amount recovered",
         "LRPY",
         "RecoveryView → on_recovery_received()",
         "post_recovery()"),

        ("Insurance Premium",
         "1105 Bank Clearing",
         "4204 Insurance Premium Income",
         "insurance_amount parameter",
         "LFEE",
         "PENDING — caller not yet wired into RepaymentView",
         "post_insurance_premium()"),

        ("MoMo Levy",
         "2109 MoMo Levy Payable",
         "1102/1103 MoMo Float",
         "Levy per MoMo Levy Act 25/2024",
         "LFEE",
         "PENDING — wire when Konse Konse gateway is live",
         "post_momo_levy()"),
    ]

    add_table(doc,
              ["Event", "Dr Account", "Cr Account", "Amount Basis", "Journal", "Trigger", "Python Method"],
              rows,
              col_widths=[3.2, 3.0, 3.2, 3.5, 1.4, 4.5, 3.5],
              font_size=7.5)

    doc.add_paragraph()
    body(doc,
         "Rows marked PENDING in the Trigger column have the Python method implemented in "
         "odoo_client.py but are not yet called from any Django view or management command. "
         "These must be wired before go-live.",
         italic=True, size=9)

    page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# Section 4 — Odoo Journals
# ═══════════════════════════════════════════════════════════════════════════════

def build_section4(doc):
    heading(doc, "SECTION 4 — Odoo Journals", 1)
    section_rule(doc)
    body(doc, (
        "Five dedicated journals were created in Odoo for the LMS integration. Standard Odoo "
        "journals (INV, BILL, MISC, BNK1, CSH1, EXCH, CABA) remain in place for other "
        "accounting operations and are not used by the LMS bridge."
    ))
    doc.add_paragraph()

    rows = [
        ("LDIS", "8",  "Loan Disbursement Journal",   "general",
         "LDIS/YYYY/MM/NNNN",
         "1111 Loan Receivable S1",
         "Records every loan disbursement. Dr 1111 / Cr 1105 per disbursement."),
        ("LRPY", "9",  "Loan Repayment Journal",       "bank",
         "LRPY/YYYY/NNNNN",
         "1105 Bank Clearing",
         "Records all cash repayments — principal, interest, and recovery postings."),
        ("LINT", "10", "Interest Income Journal",      "general",
         "LINT/YYYY/MM/NNNN",
         "1120 Interest Receivable",
         "Month-end interest accrual entries. Dr 1120 / Cr 4101/4102/4103."),
        ("LFEE", "11", "Penalty & Fees Journal",       "general",
         "LFEE/YYYY/MM/NNNN",
         "1121 Penalty Fees Receivable",
         "Origination fees, VAT on fees, penalties, insurance premiums, write-offs."),
        ("LPROV", "12", "Loan Loss Provision Journal", "general",
         "LPROV/YYYY/MM/NNNN",
         "5101 Provision Expense S1",
         "Monthly ECL provisioning and IFRS 9 stage-transfer provision entries."),
    ]
    add_table(doc,
              ["Code", "Odoo ID", "Name", "Type", "Sequence Format", "Default Account", "Purpose"],
              rows,
              col_widths=[1.2, 1.2, 4.2, 1.5, 3.5, 4.0, 5.5],
              font_size=8)

    doc.add_paragraph()
    body(doc, (
        "Journal codes are stored in .env / .env.production as ODOO_JOURNAL_DISBURSEMENT, "
        "ODOO_JOURNAL_REPAYMENT, ODOO_JOURNAL_INTEREST, ODOO_JOURNAL_FEES, and "
        "ODOO_JOURNAL_PROVISION. If journals are renamed in Odoo, the corresponding "
        ".env variable must be updated before restarting the LMS backend."
    ), size=9, italic=True)

    page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# Section 5 — Bridge Architecture
# ═══════════════════════════════════════════════════════════════════════════════

def build_section5(doc):
    heading(doc, "SECTION 5 — LMS–Odoo Bridge Architecture", 1)
    section_rule(doc)

    heading(doc, "5.1  Key Files", 2)
    rows = [
        ("apps/accounting/odoo_client.py",       "Core XML-RPC client. OdooLMSClient class with all post_*() methods. Singleton via get_odoo_client()."),
        ("apps/accounting/on_loan_approved.py",  "Disbursement hook: syncs borrower → res.partner, posts disbursement journal, persists odoo_partner_id and odoo_disbursement_move_id on the Loan record."),
        ("apps/accounting/on_payment_received.py","Repayment hook: splits principal/interest using EIR, posts LRPY entry, posts penalty if applicable."),
        ("apps/accounting/on_loan_written_off.py","Write-off hook: posts Dr 1204 / Cr 1113. Called from WriteOffLoanView."),
        ("apps/accounting/on_loan_stage_changed.py","Stage-change hook: posts provision transfer entry. Called from update_loan_statuses command."),
        ("apps/accounting/on_repayment_due.py",  "Repayment-due hook: creates an Odoo vendor bill for upcoming instalments. Called by sync_odoo_repayment_invoices cron command."),
        ("apps/accounting/konse_konse_client.py","Konse Konse HTTP API client with HMAC-SHA256 signing and retry/backoff."),
        ("apps/accounting/konse_events.py",      "Routes Konse Konse payment events to the correct Odoo journal posting method."),
        ("apps/accounting/momo_reconcile.py",    "MoMo webhook reconciliation and MoMo levy posting logic."),
        ("config/config_loader.py",              "Multi-environment .env loader. Selects .env.local / .env.production based on APP_ENV."),
        (".env.local",                           "Development environment variables (WSL). Loaded when APP_ENV=local (default)."),
        (".env.production",                      "Production environment template. Fill all CHANGE_ME values before deploying."),
    ]
    add_table(doc, ["File (relative to backend/)", "Purpose"], rows,
              col_widths=[5.5, 10.0], font_size=8)

    doc.add_paragraph()
    heading(doc, "5.2  Authentication Flow", 2)
    body(doc, "Every XML-RPC session follows this sequence:")
    steps = [
        "1.  OdooLMSClient.__init__() reads ODOO_URL, ODOO_DB, ODOO_USER, ODOO_PASSWORD from environment.",
        "2.  On first API call, _authenticate() is invoked: POST to /xmlrpc/2/common → common.authenticate() → returns integer UID.",
        "3.  Subsequent calls use _models proxy (ServerProxy at /xmlrpc/2/object) with UID + password on every call.",
        "4.  If authentication fails (bad credentials, Odoo down), OdooConnectionError is raised and logged. The LMS operation already committed to its own DB is not rolled back.",
        "5.  get_odoo_client() returns a module-level singleton — the client is created once per Django process and reused for all subsequent events.",
    ]
    for s in steps:
        bullet(doc, s)

    doc.add_paragraph()
    heading(doc, "5.3  Service Account", 2)
    rows2 = [
        ("Odoo login",      "lms_bridge"),
        ("Odoo UID",        "6"),
        ("Display name",    "LMS Bridge Service"),
        ("Permission group","Full Accounting Features (group id 24)"),
        ("Email",           "lms_bridge@intzam.zm"),
        ("Created",         "17 March 2026 via XML-RPC (not via Odoo UI)"),
        (".env variable",   "ODOO_USER=lms_bridge  /  ODOO_PASSWORD=<stored in .env.local and .env.production>"),
    ]
    add_table(doc, ["Property", "Value"], rows2, col_widths=[5, 10.5], font_size=9)

    doc.add_paragraph()
    heading(doc, "5.4  Error Handling Policy", 2)
    body(doc, (
        "All Odoo posting calls are wrapped in try/except inside the event hook files "
        "(on_loan_approved.py, on_payment_received.py, etc.). Errors are logged at ERROR "
        "level using Django's standard logging framework and do NOT propagate back to the "
        "view — the HTTP response to the frontend is always returned based on the LMS "
        "database outcome, not the Odoo outcome."
    ))
    body(doc, (
        "If ODOO_ENABLED=false is set in the environment, all XML-RPC calls are silently "
        "skipped. This is useful for offline development or maintenance windows."
    ))

    heading(doc, "5.5  Idempotency", 2)
    body(doc, (
        "The lms.loan.event model in Odoo serves as an idempotency ledger. Each posting "
        "method checks for an existing event record with the same lms_reference before "
        "creating a new journal entry. Duplicate events (e.g. from a retry after a network "
        "error) raise OdooDuplicateError, which is caught and handled gracefully — the "
        "existing journal entry ID is fetched and stored without creating a second entry."
    ))

    heading(doc, "5.6  Running the Connection Test", 2)
    body(doc, "Run the following to verify all 7 integration checkpoints:", size=9)
    p = doc.add_paragraph()
    run = p.add_run(
        "source /home/mufwaya/lms-venv/bin/activate\n"
        "cd /mnt/f/LMS/backend\n"
        "python test_odoo_connection.py"
    )
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    p.paragraph_format.left_indent = Cm(0.5)

    page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# Section 6 — Cron Job Schedule
# ═══════════════════════════════════════════════════════════════════════════════

def build_section6(doc):
    heading(doc, "SECTION 6 — Cron Job Schedule", 1)
    section_rule(doc)
    body(doc, (
        "Five cron jobs have been installed in the crontab of user mufwaya on the "
        "development (WSL) machine. These must be re-installed on the production server "
        "under the appropriate service user. All jobs use the Linux venv Python binary "
        "and log to ~/logs/lms/. Logs rotate manually — consider adding logrotate "
        "configuration in production."
    ))
    doc.add_paragraph()

    rows = [
        ("Daily at 01:00 CAT",
         "update_loan_statuses",
         "update_loan_statuses.log",
         "Scans all active loans, applies IFRS 9 stage transitions (S1→S2→S3) based on "
         "days overdue thresholds. Calls on_loan_stage_changed() to post provision "
         "transfer entries to Odoo for any stage-changed loans."),

        ("Daily at 08:00 CAT",
         "sync_odoo_repayment_invoices",
         "odoo_repayment_sync.log",
         "Creates Odoo vendor bill / repayment invoice for each instalment due within "
         "the next 24 hours. Supports --dry-run and --loan LOAN_NUMBER flags for "
         "manual testing. Skips instalments not yet due."),

        ("1st of month at 07:00 CAT",
         "sync_odoo_ecl_provision",
         "odoo_ecl_provision.log",
         "Calculates ECL (Expected Credit Loss) provision amounts by IFRS 9 stage "
         "and posts Dr 5101/5102/5103 / Cr 1201/1202/1203 to the LPROV journal. "
         "Run after update_loan_statuses has updated all stages."),

        ("Last day of month at 23:55 CAT\n(days 28–31 check)",
         "sync_odoo_interest_accrual",
         "odoo_interest_accrual.log",
         "Posts month-end EIR interest accrual: Dr 1120 Interest Receivable / "
         "Cr 4101/4102/4103 Interest Income by IFRS 9 stage. Cron fires on days 28–31 "
         "and the command itself verifies it is the last day of the month."),

        ("Every minute  (* * * * *)",
         "poll_konse_transactions",
         "konse_poll.log",
         "Polls the Konse Konse (*543#) gateway API for pending USSD payment "
         "transactions. Routes confirmed payments through konse_events.py to post "
         "the appropriate Odoo journal entry. Currently in sandbox mode — "
         "no live calls until KONSE_API_KEY and KONSE_SECRET are set."),
    ]

    tbl = doc.add_table(rows=1 + len(rows), cols=4)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_header_row(tbl, ["Schedule", "Command", "Log File", "What It Does"],
                   col_widths=[3.5, 4.5, 4.0, 9.5])
    for i, (sched, cmd, log, desc) in enumerate(rows):
        row = tbl.rows[i + 1]
        bg = 'F4F6F8' if i % 2 == 1 else 'FFFFFF'
        for cell in row.cells:
            set_cell_bg(cell, bg)
        cell_para(row.cells[0], sched, font_size=8)
        cell_para(row.cells[1], cmd, bold=True, font_size=8,
                  color=ACCENT_BLUE)
        cell_para(row.cells[2], log, font_size=8,
                  color=RGBColor(0x5D, 0x6D, 0x7E))
        cell_para(row.cells[3], desc, font_size=8)

    doc.add_paragraph()
    heading(doc, "6.1  Cron Job Dependencies", 2)
    body(doc, "The following execution order must be respected on the 1st of each month:")
    steps = [
        "1.  01:00  update_loan_statuses  — must run first to update IFRS 9 stages.",
        "2.  07:00  sync_odoo_ecl_provision  — uses the stage assignments from step 1.",
        "3.  23:55  sync_odoo_interest_accrual  — independent; can run any time before midnight.",
    ]
    for s in steps:
        bullet(doc, s)

    doc.add_paragraph()
    heading(doc, "6.2  Installing Cron Jobs on Production", 2)
    p = doc.add_paragraph()
    run = p.add_run(
        "# Install as the LMS service user on the production server:\n"
        "VENV_PYTHON=/opt/lms/venv/bin/python\n"
        "MANAGE=/opt/lms/backend/manage.py\n"
        "LOGDIR=/var/log/lms\n\n"
        "# Paste into crontab -e:\n"
        "0  1 * * *     $VENV_PYTHON $MANAGE update_loan_statuses          >> $LOGDIR/update_loan_statuses.log 2>&1\n"
        "0  8 * * *     $VENV_PYTHON $MANAGE sync_odoo_repayment_invoices  >> $LOGDIR/odoo_repayment_sync.log  2>&1\n"
        "0  7 1 * *     $VENV_PYTHON $MANAGE sync_odoo_ecl_provision       >> $LOGDIR/odoo_ecl_provision.log   2>&1\n"
        "55 23 28-31 * * $VENV_PYTHON $MANAGE sync_odoo_interest_accrual   >> $LOGDIR/odoo_interest_accrual.log 2>&1\n"
        "* * * * *      $VENV_PYTHON $MANAGE poll_konse_transactions       >> $LOGDIR/konse_poll.log           2>&1\n"
    )
    run.font.name = 'Courier New'
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    p.paragraph_format.left_indent = Cm(0.5)

    page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# Section 7 — Konse Konse Gateway
# ═══════════════════════════════════════════════════════════════════════════════

def build_section7(doc):
    heading(doc, "SECTION 7 — Konse Konse (*543#) Payment Gateway", 1)
    section_rule(doc)

    heading(doc, "7.1  Integration Status", 2)
    rows = [
        ("USSD shortcode",       "*543#  (configured in KONSE_SHORTCODE env var)"),
        ("Integration mode",     "SANDBOX — all API calls go to sandbox.konsepay.com"),
        ("Code status",          "COMPLETE — all gateway code built and tested"),
        ("Live credentials",     "PENDING — KONSE_API_KEY and KONSE_SECRET not yet obtained from Konse Konse"),
        ("Callback URL",         "PENDING — requires a public HTTPS URL (ngrok in dev; production domain in prod)"),
        ("Odoo COA accounts",    "1107 KK Gateway Float, 1108 KK Disbursement Clearing, 2111 KK Fees Payable, 5223 KK Gateway Fees — all created in Odoo"),
        ("Test suite",           "backend/test_konse_odoo_flow.py — 7-case integration test suite"),
    ]
    add_table(doc, ["Item", "Status / Value"], rows, col_widths=[5, 11], font_size=9)

    doc.add_paragraph()
    heading(doc, "7.2  Payment Flow", 2)
    body(doc, "The end-to-end Konse Konse payment flow is as follows:")
    steps = [
        "1.  Customer dials *543# on their mobile phone (MTN, Airtel, or Zamtel).",
        "2.  USSD menu prompts for Loan ID and amount.",
        "3.  Customer confirms; Konse Konse gateway processes the MoMo debit.",
        "4.  Konse Konse sends a webhook POST to /api/accounting/konse-webhook/ (Django endpoint).",
        "5.  KonseWebhookView verifies the HMAC-SHA256 signature using KONSE_WEBHOOK_SECRET.",
        "6.  A KonseTransaction record is saved to the LMS database (status=pending).",
        "7.  The poll_konse_transactions cron (every minute) picks up pending transactions.",
        "8.  konse_events.py routes the event type (disbursement / repayment / fee / agent_collection) to the correct Odoo posting method.",
        "9.  The appropriate journal entry is created in Odoo (LRPY or LDIS journal).",
        "10. KonseTransaction status is updated to processed or failed.",
    ]
    for s in steps:
        bullet(doc, s)

    doc.add_paragraph()
    heading(doc, "7.3  Key Files", 2)
    rows2 = [
        ("apps/accounting/konse_konse_client.py",
         "HTTP API client with HMAC-SHA256 request signing, retry/backoff, sandbox flag, and idempotent transaction reference."),
        ("apps/accounting/konse_events.py",
         "Event router: maps KonseTransaction.event_type to the correct odoo_client.post_*() method."),
        ("apps/loans/models.py (KonseTransaction)",
         "Django model storing each gateway transaction: reference, amount, event_type, status, timestamps."),
        ("apps/loans/migrations/0004_konse_transaction.py",
         "Django migration that adds the KonseTransaction table."),
        ("apps/core/management/commands/poll_konse_transactions.py",
         "Management command polled every minute by cron. Processes all KonseTransaction records in status=pending."),
        ("POST /api/accounting/konse-webhook/",
         "Webhook endpoint. Validates HMAC signature, saves incoming event as KonseTransaction(status=pending)."),
    ]
    add_table(doc, ["File / Endpoint", "Purpose"], rows2,
              col_widths=[6.5, 9.5], font_size=8)

    doc.add_paragraph()
    pending(doc, "Obtain KONSE_API_KEY and KONSE_SECRET from Konse Konse. Set KONSE_SANDBOX=false and update KONSE_CALLBACK_URL to the production HTTPS domain before go-live.")

    page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# Section 8 — Deployment & Environment Setup
# ═══════════════════════════════════════════════════════════════════════════════

def build_section8(doc):
    heading(doc, "SECTION 8 — Deployment & Environment Setup", 1)
    section_rule(doc)

    heading(doc, "8.1  Current Development Environment (WSL)", 2)
    rows = [
        ("Host OS",            "Windows 11 with WSL2 (Ubuntu)"),
        ("WSL distro",         "Ubuntu 22.04 LTS"),
        ("Python venv (WSL)",  "/home/mufwaya/lms-venv/  (Python 3.12)"),
        ("LMS source",         "/mnt/f/LMS/"),
        ("Django backend",     "/mnt/f/LMS/backend/"),
        ("Odoo 17 binary",     "/opt/odoo17/odoo-bin"),
        ("Odoo config",        "/etc/odoo17.conf"),
        ("Odoo custom addons", "/opt/odoo17/custom_addons/lms_bridge/"),
        ("Odoo data directory","/home/mufwaya/.local/share/Odoo/"),
        ("PostgreSQL",         "localhost:5432  (odoo_user / odoo_pass_2025)"),
        ("Odoo DB (dev)",      "odoo_lms_test"),
        ("LMS DB (dev)",       "SQLite — backend/db.sqlite3"),
        ("Start Odoo",         "bash ~/start_odoo.sh   (PID stored at /tmp/odoo17.pid)"),
        ("Stop Odoo",          "kill $(cat /tmp/odoo17.pid)"),
        ("Start Django",       "source /home/mufwaya/lms-venv/bin/activate && cd /mnt/f/LMS/backend && python manage.py runserver"),
        ("Cron log directory", "/home/mufwaya/logs/lms/"),
    ]
    add_table(doc, ["Item", "Value"], rows, col_widths=[4.5, 11], font_size=9)

    doc.add_paragraph()
    heading(doc, "8.2  Production Server Requirements", 2)
    pending(doc, "Production server not yet provisioned. Fill in the details below once a server is allocated.")
    doc.add_paragraph()
    rows2 = [
        ("Provider",         "[PENDING — e.g. Hetzner, AWS, DigitalOcean]"),
        ("OS",               "[PENDING — Ubuntu 22.04 LTS recommended]"),
        ("Spec",             "[PENDING — minimum 2 vCPU / 4 GB RAM / 40 GB SSD for Odoo + LMS]"),
        ("Public IP / Domain","[PENDING — e.g. api.intzam.co.zm]"),
        ("Odoo URL",         "[PENDING — http://127.0.0.1:8069 if co-located, or separate host]"),
        ("PostgreSQL version","PostgreSQL 16 recommended"),
        ("Reverse proxy",    "[PENDING — Nginx recommended in front of Django (Gunicorn) and Odoo]"),
        ("SSL / HTTPS",      "[PENDING — Let's Encrypt or commercial certificate]"),
        ("Odoo systemd unit","odoo17.service — use the unit file in the Odoo 17 install guide"),
        ("Django systemd unit","[PENDING — Gunicorn + systemd unit, or Docker container]"),
    ]
    add_table(doc, ["Item", "Value"], rows2, col_widths=[4.5, 11], font_size=9)

    doc.add_paragraph()
    heading(doc, "8.3  Environment Variables (.env structure)", 2)
    body(doc, (
        "The LMS backend uses a config_loader.py that selects the correct .env file based on "
        "the APP_ENV environment variable. Never commit .env files with real secrets."
    ))
    rows3 = [
        (".env.local",       "APP_ENV=local",       "WSL development. Contains dev Odoo credentials.  DO NOT commit."),
        (".env.production",  "APP_ENV=production",  "Production template — all CHANGE_ME values must be filled.  DO NOT commit."),
        (".env.example",     "(any)",               "Safe-to-commit reference showing every variable name with no values."),
    ]
    add_table(doc, ["File", "APP_ENV value", "Purpose"], rows3,
              col_widths=[3.5, 3.5, 9], font_size=9)

    doc.add_paragraph()
    heading(doc, "8.4  Critical .env Variables for the Odoo Bridge", 2)
    rows4 = [
        ("ODOO_URL",          "Full URL to the Odoo server, e.g. http://127.0.0.1:8069"),
        ("ODOO_DB",           "Odoo database name (odoo_lms_test or odoo_lms_production)"),
        ("ODOO_USER",         "Service account login — always lms_bridge, never admin"),
        ("ODOO_PASSWORD",     "Service account password — strong random string"),
        ("ODOO_ENABLED",      "true / false — set false to disable all Odoo calls (maintenance mode)"),
        ("ODOO_JOURNAL_*",    "Five variables mapping journal purpose to journal code (LDIS, LRPY, LINT, LFEE, LPROV)"),
        ("ODOO_ACCOUNT_*",    "Nineteen variables mapping account purpose to account code (1105, 1111, 4101, etc.)"),
        ("KONSE_API_KEY",     "Konse Konse API key — obtain from the Konse Konse merchant portal"),
        ("KONSE_SECRET",      "Konse Konse signing secret — used to verify webhook HMAC-SHA256 signatures"),
        ("KONSE_SANDBOX",     "true in dev/staging, false in production"),
        ("KONSE_CALLBACK_URL","Public HTTPS URL for the Konse Konse webhook endpoint"),
    ]
    add_table(doc, ["Variable", "Description"], rows4,
              col_widths=[4.5, 11.5], font_size=9)

    doc.add_paragraph()
    heading(doc, "8.5  Backup & Recovery", 2)
    rows5 = [
        ("Odoo database",  "pg_dump via export_odoo.sh",       "Retained 30 days",  "bash /mnt/f/LMS/export_odoo.sh  →  ~/odoo_exports/odoo_export_YYYYMMDD.zip"),
        ("Odoo filestore", "rsync — included in export_odoo.sh", "Retained 30 days","Bundled inside the dated .zip archive alongside the SQL dump"),
        ("LMS database (dev)", "SQLite file backup",           "Manual",            "/mnt/f/LMS/backend/db.sqlite3 — back up the file directly"),
        ("LMS database (prod)","pg_dump (separate cron)",      "Retained 30 days",  "[PENDING — set up daily pg_dump cron on the production server]"),
        ("lms_bridge module","Included in export_odoo.sh",     "With each export",  "custom_addons/lms_bridge/ — also stored in source control"),
    ]
    add_table(doc,
              ["What", "Method", "Retention", "Location / Command"],
              rows5, col_widths=[3.5, 4, 2.5, 7], font_size=8)

    doc.add_paragraph()
    heading(doc, "8.6  Deploying to Production (Step-by-Step)", 2)
    steps = [
        "1.  Provision the production server (Ubuntu 22.04, min 4 GB RAM).",
        "2.  Install PostgreSQL 16, Odoo 17 Community, Python 3.12, and Nginx.",
        "3.  On the dev machine, run:  bash /mnt/f/LMS/export_odoo.sh",
        "       This produces ~/odoo_exports/odoo_export_YYYYMMDD.zip  (DB dump + filestore + lms_bridge + config).",
        "4.  Transfer the archive:  scp ~/odoo_exports/odoo_export_*.zip user@PROD_SERVER:/opt/odoo_imports/",
        "5.  Transfer the import script:  scp /mnt/f/LMS/import_odoo.sh user@PROD_SERVER:/opt/",
        "6.  On the production server, run:  bash /opt/import_odoo.sh --archive /opt/odoo_imports/odoo_export_*.zip --target-db odoo_lms_production",
        "       The script: extracts the archive, restores the DB, copies the filestore, installs lms_bridge, and updates the config.",
        "7.  Edit /etc/odoo17.conf: set db_name, db_password, admin_passwd, workers, logfile, data_dir.",
        "8.  Copy backend/.env.production to /opt/lms/backend/.env.production — fill ALL CHANGE_ME values.",
        "9.  Set permissions:  chmod 600 /opt/lms/backend/.env.production",
        "10. Create the Odoo service account in production:  run the lms_bridge account creation script or create manually in Odoo UI.",
        "11. Start Odoo:  sudo systemctl start odoo17",
        "12. Run the connection test:  APP_ENV=production python manage.py shell -c \"from apps.accounting.odoo_client import get_odoo_client; print(get_odoo_client().ping())\"",
        "13. Install the 5 cron jobs (see Section 6.2) under the LMS service user.",
        "14. Configure Nginx as a reverse proxy for both Django (port 8000) and Odoo (port 8069).",
        "15. Install SSL certificate and enable HTTPS; then uncomment the SECURE_SSL_REDIRECT settings in .env.production.",
    ]
    for s in steps:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(s)
        run.font.size = Pt(9)

    page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# Section 9 — Handover Checklist
# ═══════════════════════════════════════════════════════════════════════════════

def build_section9(doc):
    heading(doc, "SECTION 9 — Handover Checklist", 1)
    section_rule(doc)
    body(doc, (
        "The IT team should work through this checklist in order after the production "
        "server has been provisioned and the system deployed. Items marked [DEV VERIFIED] "
        "have been confirmed working on the development environment as of 17 March 2026."
    ))
    doc.add_paragraph()

    groups = [
        ("A. Infrastructure", [
            ("Production server provisioned (OS, PostgreSQL, Python 3.12, Nginx)", False),
            ("Odoo 17 Community installed at /opt/odoo17", False),
            ("LMS backend deployed to /opt/lms/backend", False),
            ("Systemd units created for Odoo (odoo17.service) and Django (lms.service)", False),
            ("SSL certificate installed and HTTPS active on production domain", False),
        ]),
        ("B. Odoo Configuration", [
            ("Odoo runs at localhost:8069 and health endpoint returns {\"status\": \"pass\"}", True),
            ("Database: odoo_lms_production created and lms_bridge module installed", False),
            ("All 5 LMS journals present (LDIS, LRPY, LINT, LFEE, LPROV)", True),
            ("All 35 LMS COA accounts present (run test_odoo_connection.py)", True),
            ("Company: IntZam Microfinance Ltd, Country: Zambia, Currency: ZMW", True),
            ("lms_bridge version 17.0.1.0.0 shows as Installed in Odoo Apps", True),
        ]),
        ("C. Service Account", [
            ("lms_bridge user exists in Odoo with Full Accounting Features permission", True),
            ("ODOO_USER=lms_bridge and ODOO_PASSWORD set correctly in .env.production", False),
            ("admin account NOT used in .env.production ODOO_USER variable", False),
            ("Connection test passes all 7 steps on production: python test_odoo_connection.py", False),
        ]),
        ("D. Event Hooks", [
            ("Live disbursement test: odoo_partner_id and odoo_disbursement_move_id populated on Loan record", True),
            ("Live disbursement test: LDIS journal entry visible in Odoo Accounting → Journal Entries", True),
            ("Live repayment test: LRPY entry with correct Dr/Cr principal/interest split", True),
            ("Write-off hook: WriteOffLoanView triggers on_loan_written_off → post_writeoff() call confirmed", True),
            ("Recovery hook: RecoveryView triggers on_recovery_received → post_recovery() call confirmed", True),
            ("Stage-change hook: update_loan_statuses triggers post_stage_transfer() for stage transitions", True),
            ("Origination fee and VAT: automatically posted alongside disbursement when processing_fee > 0", True),
            ("Insurance premium posting: PENDING — not yet wired into RepaymentView", False),
        ]),
        ("E. Cron Jobs", [
            ("All 5 cron jobs installed in production crontab", False),
            ("/var/log/lms/ directory exists with correct permissions", False),
            ("update_loan_statuses dry-run produces no errors", True),
            ("sync_odoo_repayment_invoices --dry-run produces no errors", True),
            ("sync_odoo_ecl_provision dry-run produces no errors", False),
            ("sync_odoo_interest_accrual dry-run produces no errors", False),
        ]),
        ("F. Konse Konse Gateway", [
            ("KONSE_API_KEY and KONSE_SECRET obtained from Konse Konse merchant portal", False),
            ("KONSE_SANDBOX=false set in .env.production", False),
            ("KONSE_CALLBACK_URL set to production HTTPS domain", False),
            ("Webhook endpoint POST /api/accounting/konse-webhook/ reachable from the internet", False),
            ("test_konse_odoo_flow.py passes all 7 cases with live credentials", False),
        ]),
        ("G. Backup & Recovery", [
            ("export_odoo.sh tested and produces a valid .zip on production server", False),
            ("Daily pg_dump cron job installed for the LMS PostgreSQL database", False),
            ("Backup files retained for 30 days (verify storage location and retention policy)", False),
            ("import_odoo.sh --dry-run tested on a staging copy of the production archive", False),
        ]),
        ("H. Security", [
            (".env.production has chmod 600 permissions", False),
            ("DEBUG=False in .env.production", False),
            ("SECURE_SSL_REDIRECT=True uncommented once HTTPS confirmed working", False),
            ("PostgreSQL: public access revoked on odoo_lms_production database", False),
            ("Odoo admin_passwd changed from test value in /etc/odoo17.conf", False),
        ]),
    ]

    for group_name, items in groups:
        heading(doc, group_name, 2)
        for text, verified in items:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.3)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            if verified:
                tick = "☑"
                status = " [DEV VERIFIED]"
                status_color = GREEN_DONE
            else:
                tick = "☐"
                status = " [PRODUCTION — TO DO]"
                status_color = RED_PENDING
            run1 = p.add_run(f"{tick}  {text}")
            run1.font.size = Pt(9)
            run2 = p.add_run(status)
            run2.font.size = Pt(8)
            run2.font.color.rgb = status_color
            run2.bold = True
        doc.add_paragraph()

    page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# Section 10 — Known Gaps & Next Steps
# ═══════════════════════════════════════════════════════════════════════════════

def build_section10(doc):
    heading(doc, "SECTION 10 — Known Gaps & Next Steps", 1)
    section_rule(doc)
    body(doc, (
        "The following items are known gaps that must be resolved before the system "
        "can be considered fully operational in production. Items are prioritised P1 "
        "(blocking go-live) through P3 (enhancement)."
    ))
    doc.add_paragraph()

    rows = [
        ("P1", "Production server not provisioned",
         "No server has been allocated or configured. All integration testing has been "
         "done on a local WSL development machine.",
         "Provision a production server (min 2 vCPU / 4 GB RAM / 40 GB SSD). "
         "Follow the 15-step deployment checklist in Section 8.6."),

        ("P1", ".env.production still a template",
         "All CHANGE_ME placeholder values in .env.production must be filled before "
         "the backend can start on the production server.",
         "Fill in: SECRET_KEY, ALLOWED_HOSTS, DATABASE_URL, ODOO_URL, ODOO_DB, "
         "ODOO_USER, ODOO_PASSWORD, KONSE_API_KEY, KONSE_SECRET, EMAIL credentials."),

        ("P1", "Konse Konse live credentials not obtained",
         "KONSE_API_KEY and KONSE_SECRET are empty. The gateway code is complete and "
         "sandbox-tested, but no live transactions can be processed.",
         "Contact Konse Konse / Zamtel to obtain sandbox-to-live API credentials. "
         "Set KONSE_SANDBOX=false and KONSE_CALLBACK_URL in .env.production."),

        ("P2", "Insurance premium posting not wired",
         "post_insurance_premium() exists in odoo_client.py and posts Dr 1105 / Cr 4204, "
         "but no Django view or cron job calls it. Insurance collections are not reflected "
         "in the Odoo GL.",
         "Add an insurance_amount optional parameter to RepaymentView. When non-zero, "
         "call on_payment_received() with the insurance amount and trigger "
         "post_insurance_premium()."),

        ("P2", "MoMo Levy not wired",
         "post_momo_levy() exists in odoo_client.py to post the MoMo Levy under the "
         "MoMo Levy Act 25/2024, but it is not called from any view. Levy applies to all "
         "MoMo-based disbursements and repayments.",
         "Wire post_momo_levy() into konse_events.py when the Konse Konse gateway goes "
         "live. The levy rate and threshold are defined in the Act."),

        ("P2", "Production cron jobs not installed",
         "Cron jobs are currently installed on the WSL dev machine only. They must be "
         "re-installed on the production server under the correct service user.",
         "Install all 5 cron jobs as shown in Section 6.2. Create /var/log/lms/ with "
         "appropriate permissions and configure logrotate."),

        ("P2", "LMS production database backup not configured",
         "A 30-day retention policy is defined but no pg_dump cron has been created for "
         "the production LMS PostgreSQL database.",
         "Install a daily pg_dump cron job for the LMS database alongside the Odoo "
         "export_odoo.sh backup. Store backups offsite (S3 or equivalent)."),

        ("P3", "Odoo not registered as a systemd service on dev",
         "Odoo in the WSL dev environment is started manually via bash ~/start_odoo.sh. "
         "It does not auto-start on WSL reboot.",
         "On the production server, Odoo must be run as a systemd service (odoo17.service) "
         "so it restarts automatically after reboots or crashes."),

        ("P3", "Swagger / OpenAPI documentation absent",
         "The LMS REST API has no auto-generated documentation. New developers must "
         "read views.py and serializers.py to understand the API surface.",
         "Add drf-spectacular to the Django project and expose /api/schema/ and "
         "/api/docs/ endpoints."),

        ("P3", "Odoo 19 upgrade planning",
         "Odoo 17 LTS is supported until 2027. Odoo 19 was released October 2025. "
         "The upgrade path is 17 → 18 → 19 (no version skipping). lms_bridge will need "
         "API changes for each major version.",
         "Plan migration for Q3/Q4 2027. Detailed breaking-change analysis is "
         "documented in the Odoo Upgrade Plan section of ODOO_INTEGRATION_STATUS.md."),
    ]

    add_table(doc,
              ["Priority", "Gap", "Impact", "Recommended Action"],
              rows,
              col_widths=[1.5, 4.0, 6.0, 7.0],
              font_size=8)

    doc.add_paragraph()
    doc.add_paragraph()
    body(doc, "— End of Document —", bold=True,
         color=BRAND_BLUE)
    p = doc.add_paragraph()
    run = p.add_run(
        f"IntZam Microfinance Ltd  ·  LMS–Odoo IT Handover  ·  {PREPARED_DATE}  ·  CONFIDENTIAL"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ═══════════════════════════════════════════════════════════════════════════════
# Main
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)

    # ── Default paragraph font ────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # ── Heading styles ────────────────────────────────────────────────────────
    for lvl, sz, bold in [(1, 13, True), (2, 11, True)]:
        h = doc.styles[f'Heading {lvl}']
        h.font.name = 'Calibri'
        h.font.size = Pt(sz)
        h.font.bold = bold
        h.font.color.rgb = BRAND_BLUE
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after  = Pt(4)

    # ── Build all sections ────────────────────────────────────────────────────
    build_cover(doc)
    build_section1(doc)
    build_section2(doc)
    build_section3(doc)
    build_section4(doc)
    build_section5(doc)
    build_section6(doc)
    build_section7(doc)
    build_section8(doc)
    build_section9(doc)
    build_section10(doc)

    doc.save(OUTPUT_PATH)
    print(f"Document saved → {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
